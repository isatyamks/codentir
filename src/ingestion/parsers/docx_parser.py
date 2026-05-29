from pathlib import Path
from typing import Optional

from src.ingestion.parsers.base_parser import BaseParser, ParsedDocument


class DOCXParser(BaseParser):
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        self._check_dependencies()
    
    def _check_dependencies(self):
        try:
            import docx
            self.has_docx = True
        except ImportError:
            raise ImportError(
                "python-docx is not installed. "
                "Install with: pip install python-docx"
            )
    
    def parse(self, file_path: Path) -> ParsedDocument:
        self.validate_file(file_path)
        
        import docx
        
        doc = docx.Document(file_path)
        
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text)
        
        content = "\n\n".join(paragraphs_text)
        
        metadata = self.extract_basic_metadata(file_path)
        metadata.update({
            "parser": "DOCXParser",
            "paragraph_count": len(doc.paragraphs),
            "section_count": len(doc.sections),
            "table_count": len(doc.tables),
            "word_count": len(content.split()),
            "char_count": len(content)
        })
        
        if doc.core_properties:
            metadata["docx_properties"] = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
            }
        
        if doc.tables:
            tables_content = self._extract_tables(doc.tables)
            if tables_content:
                content += "\n\n[TABLES]\n" + tables_content
        
        return ParsedDocument(
            file_path=file_path,
            content=content,
            file_type="docx",
            metadata=metadata
        )
    
    def _extract_tables(self, tables) -> str:
        tables_text = []
        
        for i, table in enumerate(tables):
            table_text = [f"\n[Table {i+1}]"]
            
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            
            tables_text.append("\n".join(table_text))
        
        return "\n\n".join(tables_text)
